from docx.shared import Pt

def desenhar_cabecalho(doc, pacote_cabecalho, pacote_servicos):
    
    # =========================================================
    # 1. PREENCHENDO A TABELA ORIGINAL DO TEMPLATE
    # =========================================================
    if len(doc.tables) > 0:
        # Pega a tabela que já existe no Word (aquela azul do topo direito)
        tabela_header = doc.tables[0]
        
        # A ordem exata dos campos na tabela original
        chaves_dicionario = ["Projeto", "Local", "Ticket", "Cliente", "Site", "Contato", "Telefone", "Técnico(s)", "Início", "Término"]

        # Faz um loop pelas 10 linhas e injeta o texto na Coluna 1 (direita)
        for i in range(10):
            try:
                celula_dir = tabela_header.cell(i, 1)
                texto_inserido = pacote_cabecalho.get(chaves_dicionario[i], "")
                
                # Injeta o texto na célula
                celula_dir.text = texto_inserido
                
                # Se o técnico digitou algo, formata em negrito
                if texto_inserido.strip() != "":
                    for paragrafo in celula_dir.paragraphs:
                        for run in paragrafo.runs:
                            run.bold = True
            except IndexError:
                continue
    else:
        print("⚠️ ERRO: Tabela original não encontrada no template!")

    # =========================================================
    # 2. INJEÇÃO DOS SERVICE DETAILS (Checkboxes)
    # =========================================================
    p_serv = doc.add_paragraph()
    p_serv.paragraph_format.space_before = Pt(10)
    run_titulo = p_serv.add_run("Service Details Realizados:\n")
    run_titulo.bold = True
    run_titulo.font.size = Pt(12)
    
    servicos_selecionados = [servico for servico, marcado in pacote_servicos.items() if marcado == 1]
    
    if servicos_selecionados:
        # AQUI ESTÁ A MUDANÇA: Usamos \n para jogar cada item para a linha de baixo
        p_serv.add_run("☑ " + "\n☑ ".join(servicos_selecionados))
    else:
        p_serv.add_run("Nenhum serviço marcado.")
        
    return doc