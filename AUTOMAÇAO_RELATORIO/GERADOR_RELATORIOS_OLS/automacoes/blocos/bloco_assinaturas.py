import os
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document

def forcar_bordas(tabela):
    tbl = tabela._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblBorders = OxmlElement('w:tblBorders')
    for borda in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        tag = OxmlElement(f'w:{borda}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), '000000')
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def pintar_fundo_celula(celula, cor_hex):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_hex))
    celula._tc.get_or_add_tcPr().append(shading_elm)

def desenhar_assinaturas(doc, pacote_assinaturas):
    # Espaço em branco e criação da tabela (agora com 8 linhas!)
    doc.add_paragraph().paragraph_format.space_before = Pt(24)
    tabela = doc.add_table(rows=8, cols=4)
    forcar_bordas(tabela)

    larguras = [Cm(2.0), Cm(6.5), Cm(2.0), Cm(6.5)]
    for row in tabela.rows:
        for i, cell in enumerate(row.cells):
            cell.width = larguras[i]

    # =========================================================
    # LINHA 0: CABEÇALHOS MESCLADOS (AZUL)
    # =========================================================
    celula_tec = tabela.cell(0, 0).merge(tabela.cell(0, 1))
    celula_cli = tabela.cell(0, 2).merge(tabela.cell(0, 3))

    celula_tec.text = "Técnico"
    celula_cli.text = "Representante do Cliente"

    for celula in [celula_tec, celula_cli]:
        pintar_fundo_celula(celula, "17365D")
        p = celula.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # =========================================================
    # LINHA 1: NOME
    # =========================================================
    tabela.cell(1, 0).text = "Nome"
    tabela.cell(1, 0).paragraphs[0].runs[0].bold = True
    tabela.cell(1, 1).text = "[NOME_TECNICO]"  # <--- CORRIGIDO PARA COLUNA 1 (1, 1)
    
    tabela.cell(1, 2).text = "Nome"
    tabela.cell(1, 2).paragraphs[0].runs[0].bold = True
    tabela.cell(1, 3).text = "[NOME_CLIENTE]"

    # =========================================================
    # LINHA 2: DATA
    # =========================================================
    tabela.cell(2, 0).text = "Data"
    tabela.cell(2, 0).paragraphs[0].runs[0].bold = True
    tabela.cell(2, 1).text = "[DATA]"  # <--- CORRIGIDO PARA COLUNA 1 (2, 1)
    
    tabela.cell(2, 2).text = "Data"
    tabela.cell(2, 2).paragraphs[0].runs[0].bold = True
    tabela.cell(2, 3).text = "[DATA_CLIENTE]"

    # =========================================================
    # LINHAS 3 a 6: DOCUMENTAÇÃO DO CLIENTE
    # =========================================================
    # IMO
    tabela.cell(3, 2).text = "IMO"
    tabela.cell(3, 2).paragraphs[0].runs[0].bold = True
    tabela.cell(3, 3).text = "[IMO_CLIENTE]"

    # AB
    tabela.cell(4, 2).text = "AB"
    tabela.cell(4, 2).paragraphs[0].runs[0].bold = True
    tabela.cell(4, 3).text = "[AB_CLIENTE]"

    # CIR
    tabela.cell(5, 2).text = "CIR"
    tabela.cell(5, 2).paragraphs[0].runs[0].bold = True
    tabela.cell(5, 3).text = "[CIR_CLIENTE]"

    # SDPO
    tabela.cell(6, 2).text = "SDPO"
    tabela.cell(6, 2).paragraphs[0].runs[0].bold = True
    tabela.cell(6, 3).text = "[SDPO_CLIENTE]"

    # Mesclando as células em branco do lado do técnico para um visual mais limpo
    tabela.cell(3, 0).merge(tabela.cell(6, 0))
    tabela.cell(3, 1).merge(tabela.cell(6, 1))

    # =========================================================
    # LINHA 7: ASSINATURA E INJEÇÃO DA IMAGEM
    # =========================================================
    tabela.cell(7, 0).text = "Assinatura"
    tabela.cell(7, 0).paragraphs[0].runs[0].bold = True
    
    tabela.cell(7, 2).text = "Assinatura"
    tabela.cell(7, 2).paragraphs[0].runs[0].bold = True
    
    # 1. PLANTAMOS AS DUAS ISCAS COMO REGRA PADRÃO ABSOLUTA
    tabela.cell(7, 1).text = "[ASSINATURA_TECNICO]"
    tabela.cell(7, 3).text = "[ASSINATURA_CLIENTE]" 

    # 2. VERIFICAMOS SE O TÉCNICO JÁ ASSINOU NA HORA (SOBRESCRITA)
    # Pegamos o caminho e o nome do técnico do dicionário
    caminho_imagem_tecnico = pacote_assinaturas.get("caminho_assinatura")
    nome_tec = pacote_assinaturas.get("nome_tecnico", "")
    
    # O Pulo do Gato: Só entra no IF se o arquivo existir E se o nome não estiver vazio!
    if caminho_imagem_tecnico and os.path.exists(caminho_imagem_tecnico) and nome_tec.strip() != "":
        # O técnico assinou de forma válida! Apagamos a isca e colocamos a foto.
        tabela.cell(7, 1).text = "" 
        
        celula_alvo = tabela.cell(7, 1)
        p = celula_alvo.paragraphs[0]
        run = p.add_run()
        try:
            run.add_picture(caminho_imagem_tecnico, width=Cm(4.5))
        except Exception as e:
            print(f"[ERRO] Falha ao inserir imagem do técnico: {e}")
            tabela.cell(7, 1).text = "[ERRO NA IMAGEM]"

    # =========================================================
    # NOVA TRAVA: BLOQUEIO TOTAL DE QUEBRA VIA XML
    # =========================================================
    for row in tabela.rows:
        # 1. Impede que a própria linha seja rasgada ao meio pelo Word
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), 'true')
        trPr.append(cantSplit)
        
        # 2. Força o agrupamento de todas as células
        for cell in row.cells:
            for paragrafo in cell.paragraphs:
                paragrafo.paragraph_format.keep_with_next = True

    return doc

def travar_tabela_para_nao_quebrar(tabela):
    for row in tabela.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), 'true')
        trPr.append(cantSplit)
        for cell in row.cells:
            for paragrafo in cell.paragraphs:
                paragrafo.paragraph_format.keep_with_next = True

def injetar_assinaturas_finais(caminho_docx, caminho_imagem_cliente=None, caminho_imagem_tecnico=None,
                                nome_cliente=None, data_cliente=None, imo_cliente=None,
                                ab_cliente=None, cir_cliente=None, sdpo_cliente=None,
                                nome_tecnico=None, data_tecnico=None): # <--- TÉCNICO INCLUÍDO AQUI!
    doc = Document(caminho_docx)

    # Agora o dicionário é um radar universal!
    dados_texto = {
        "nome_cliente": nome_cliente,
        "data_cliente": data_cliente,
        "imo_cliente": imo_cliente,
        "ab_cliente": ab_cliente,
        "cir_cliente": cir_cliente,
        "sdpo_cliente": sdpo_cliente,
        "nome_tecnico": nome_tecnico, # <--- Isca [NOME_TECNICO]
        "data": data_tecnico          # <--- Isca [DATA] (A chave precisa ser "data" para bater com a isca)
    }

    for tabela in doc.tables:
        for row in tabela.rows:
            for cell in row.cells:
                
                # --- CAÇADOR DE ISCAS DE TEXTO ---
                for chave, valor in dados_texto.items():
                    isca = f"[{chave.upper()}]"
                    if isca in cell.text:
                        texto_substituto = valor if valor else ""
                        cell.text = cell.text.replace(isca, texto_substituto)

                # --- CAÇADOR DE ASSINATURAS DO CLIENTE ---
                if "[ASSINATURA_CLIENTE]" in cell.text:
                    cell.text = ""
                    if caminho_imagem_cliente and os.path.exists(caminho_imagem_cliente):
                        p = cell.paragraphs[0]
                        run = p.add_run()
                        run.add_picture(caminho_imagem_cliente, width=Cm(4.5))

                # --- CAÇADOR DE ASSINATURAS DO TÉCNICO ---
                if "[ASSINATURA_TECNICO]" in cell.text:
                    cell.text = ""
                    if caminho_imagem_tecnico and os.path.exists(caminho_imagem_tecnico):
                        p = cell.paragraphs[0]
                        run = p.add_run()
                        try:
                            run.add_picture(caminho_imagem_tecnico, width=Cm(4.5))
                        except Exception as e:
                            print(f"[ERRO] Falha ao inserir imagem do técnico: {e}")
                            cell.text = "[ERRO NA IMAGEM]"

    # --- TRAVA FINAL: reaplica depois de qualquer substituição de isca ---
    for tabela in doc.tables:
        travar_tabela_para_nao_quebrar(tabela)
        
    nome_final = caminho_docx.replace(".docx", "_Relatorio_Finalizado.docx")
    doc.save(nome_final)
    return nome_final

def injetar_assinatura_tecnico(doc, caminho_imagem_tecnico=None, nome_tecnico=None, data_tecnico=None):

    dados_texto = {
        "nome_tecnico": nome_tecnico,  # Isca [NOME_TECNICO]
        "data": data_tecnico           # Isca [DATA]
    }

    for tabela in doc.tables:
        for row in tabela.rows:
            for cell in row.cells:

                # --- CAÇADOR DE ISCAS DE TEXTO ---
                for chave, valor in dados_texto.items():
                    isca = f"[{chave.upper()}]"
                    # só mexe se REALMENTE tiver valor preenchido; senão, isca fica intocada
                    if valor and isca in cell.text:
                        cell.text = cell.text.replace(isca, valor)

                # --- CAÇADOR DE ASSINATURA DO TÉCNICO ---
                # só apaga a isca se houver imagem de verdade pra colocar no lugar
                if caminho_imagem_tecnico and os.path.exists(caminho_imagem_tecnico) and "[ASSINATURA_TECNICO]" in cell.text:
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run()
                    try:
                        run.add_picture(caminho_imagem_tecnico, width=Cm(4.5))
                    except Exception as e:
                        print(f"[ERRO] Falha ao inserir imagem do técnico: {e}")
                        cell.text = "[ERRO NA IMAGEM]"

    # --- TRAVA FINAL: reaplica depois de qualquer substituição de isca ---
    for tabela in doc.tables:
        travar_tabela_para_nao_quebrar(tabela)

    return doc