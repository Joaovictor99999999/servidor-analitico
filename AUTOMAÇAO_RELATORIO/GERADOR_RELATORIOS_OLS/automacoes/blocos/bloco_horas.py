from docx.shared import Pt, Cm, RGBColor
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

def forcar_bordas(tabela):
    tbl = tabela._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblBorders = OxmlElement('w:tblBorders')
    for borda in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        tag = OxmlElement(f'w:{borda}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), '000000')
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def pintar_fundo_celula(celula, cor_hex):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_hex))
    celula._tc.get_or_add_tcPr().append(shading_elm)

def alinhar_verticalmente_centro(celula):
    """Injeta código XML para alinhar o texto exatamente no meio (vertical) da célula."""
    tcPr = celula._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def desenhar_service_log(doc, pacote_horas):
    cor_azul_ols = "17365D"
    
    # Título da Seção
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(18)
    run_titulo = p_titulo.add_run("Service Log")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    
    # =========================================================
    # TABELA 1: APONTAMENTOS DIÁRIOS
    # =========================================================
    tabela_log = doc.add_table(rows=1, cols=6)
    forcar_bordas(tabela_log)
    
    larguras_log = [Cm(1.5), Cm(3.0), Cm(2.0), Cm(2.0), Cm(2.0), Cm(6.5)]
    cabecalhos_log = ['SVC', 'Data', 'Início', 'Término', 'Total', 'Descrição']
    
    for i, nome in enumerate(cabecalhos_log):
        celula = tabela_log.cell(0, i)
        celula.width = larguras_log[i]
        pintar_fundo_celula(celula, cor_azul_ols)
        
        celula.text = nome
        p = celula.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    resumo_horas = {
        "1": {"nome": "1 - Trabalho Onshore", "total": 0.0},
        "2": {"nome": "2 - Trabalho Offshore", "total": 0.0},
        "3": {"nome": "3 - Em Stand-By (Onshore)", "total": 0.0},
        "4": {"nome": "4 - Em Stand-By (Offshore)", "total": 0.0},
        "5": {"nome": "5 - Deslocamento Terrestre", "total": 0.0},
        "6": {"nome": "6 - Deslocamento Aéreo", "total": 0.0},
        "7": {"nome": "7 - Deslocamento Aéreo (Offshore)", "total": 0.0},
        "8": {"nome": "8 - Atividade no Escritório / Warehouse", "total": 0.0}
    }

    formato_hora = "%H:%M"
    for item in pacote_horas:
        svc_str = item.get("tipo", "")
        svc_codigo = svc_str.split(" - ")[0] if " - " in svc_str else "0"
        inicio = item.get("inicio", "00:00")
        fim = item.get("fim", "00:00")
        
        horas_trabalhadas = 0.0
        try:
            t1 = datetime.strptime(inicio, formato_hora)
            t2 = datetime.strptime(fim, formato_hora)
            diff = t2 - t1
            horas_trabalhadas = diff.total_seconds() / 3600.0 
            if horas_trabalhadas < 0:
                horas_trabalhadas += 24.0
        except ValueError:
            horas_trabalhadas = 0.0

        if svc_codigo in resumo_horas:
            resumo_horas[svc_codigo]["total"] += horas_trabalhadas

        row_cells = tabela_log.add_row().cells
        h_int = int(horas_trabalhadas)
        m_int = int(round((horas_trabalhadas - h_int) * 60))
        hora_formatada = f"{h_int:02d}:{m_int:02d}"

        dados_linha = [
            svc_codigo, item.get("data", ""), inicio, fim, 
            hora_formatada, item.get("descricao", "")
        ]
        
        for i, dado in enumerate(dados_linha):
            row_cells[i].text = dado
            row_cells[i].width = larguras_log[i]
            if i in [0, 1, 2, 3, 4]: # Centraliza as colunas de dados, menos a descrição
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================
    # TABELA 2: RESUMO FINAL (COM MESCLAGEM)
    # =========================================================
    doc.add_paragraph().paragraph_format.space_before = Pt(12) # Espaço entre as tabelas

    # Cria tabela com 9 linhas (1 cabeçalho + 8 serviços) e 3 colunas
    tabela_resumo = doc.add_table(rows=9, cols=3)
    forcar_bordas(tabela_resumo)
    
    larguras_resumo = [Cm(8.0), Cm(3.0), Cm(6.0)]
    cabecalhos_resumo = ['Service Code', 'Horas', 'Total de Horas']
    
    # Pinta e formata o cabeçalho
    for i, nome in enumerate(cabecalhos_resumo):
        celula = tabela_resumo.cell(0, i)
        celula.width = larguras_resumo[i]
        pintar_fundo_celula(celula, cor_azul_ols)
        celula.text = nome
        p = celula.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Preenche os serviços e acumula o total
    soma_total_absoluto = 0.0
    codigos = ["1", "2", "3", "4", "5", "6", "7", "8"]
    
    for idx, codigo in enumerate(codigos):
        linha_atual = idx + 1
        nome = resumo_horas[codigo]["nome"]
        total = resumo_horas[codigo]["total"]
        soma_total_absoluto += total
        
        # Coluna 0 (Service Code)
        cel_codigo = tabela_resumo.cell(linha_atual, 0)
        cel_codigo.text = nome
        cel_codigo.width = larguras_resumo[0]
        alinhar_verticalmente_centro(cel_codigo)
        
        # Coluna 1 (Horas)
        cel_horas = tabela_resumo.cell(linha_atual, 1)
        
        # Converte decimal para HH:MM
        h_total = int(total)
        m_total = int(round((total - h_total) * 60))
        cel_horas.text = f"{h_total:02d}:{m_total:02d}"
        
        cel_horas.width = larguras_resumo[1]
        cel_horas.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        alinhar_verticalmente_centro(cel_horas)

    # O GRANDE TRUQUE: Mesclar todas as linhas da Coluna 2 (Total Absoluto)
    celula_topo = tabela_resumo.cell(1, 2)
    celula_base = tabela_resumo.cell(8, 2)
    celula_mesclada = celula_topo.merge(celula_base)
    
    # Insere o Totalzão e formata gigante e centralizado
    # Insere o Totalzão e formata gigante e centralizado
    h_abs = int(soma_total_absoluto)
    m_abs = int(round((soma_total_absoluto - h_abs) * 60))
    celula_mesclada.text = f"{h_abs:02d}:{m_abs:02d}"
    
    celula_mesclada.width = larguras_resumo[2]
    alinhar_verticalmente_centro(celula_mesclada)
    
    p_total = celula_mesclada.paragraphs[0]
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_total = p_total.runs[0]
    run_total.font.size = Pt(28) # Fonte gigante igual da foto

    return doc