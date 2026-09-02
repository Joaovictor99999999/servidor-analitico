import os
from io import BytesIO
from PIL import Image
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def anexar_fotos_do_bloco(doc, lista_fotos):
    """
    Motor de fotos inteligente adaptado para a Interface Gráfica.
    Recebe uma lista de dicionários contendo 'caminho' e 'legenda'.
    """
    # Se a lista estiver vazia (o usuário não adicionou fotos na interface), não faz nada
    if not lista_fotos:
        return doc
        
    print(f"📸 Anexando {len(lista_fotos)} fotos selecionadas pela interface...")
    
    # 🚨 1. A QUEBRA DE PÁGINA COM RESPIRO DE TOPO
    doc.add_page_break()
    p_topo = doc.add_paragraph()
    p_topo.paragraph_format.space_after = Pt(24) # Empurra tudo para baixo logo de cara
    
    # Agora iteramos diretamente na lista que veio da interface
    for item in lista_fotos:
        caminho_imagem = item["caminho"]
        legenda_texto = item["legenda"] or "Foto do Sistema" # Padrão caso o usuário deixe a legenda vazia
        
        try:
            img = Image.open(caminho_imagem)
            
            # A INTELIGÊNCIA: Mede os pixels originais antes de qualquer coisa
            largura, altura = img.size 
            
            img.thumbnail((1920, 1080)) 
            
            imagem_memoria = BytesIO()
            img.save(imagem_memoria, format="PNG") 
            imagem_memoria.seek(0)
            
            # 🚨 2. PARÁGRAFO DA LEGENDA (Com Margem Superior)
            p_legenda = doc.add_paragraph()
            p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # O SEGREDO AQUI: Garante que a legenda se afaste do topo ou da foto anterior
            p_legenda.paragraph_format.space_before = Pt(18) 
            
            run_legenda = p_legenda.add_run(legenda_texto)
            run_legenda.bold = True
            p_legenda.paragraph_format.keep_with_next = True 
            p_legenda.paragraph_format.space_after = Pt(6)
            
            # 2. PARÁGRAFO DA FOTO
            p_foto = doc.add_paragraph()
            p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 🌳 ÁRVORE DE DECISÃO (Paisagem vs Retrato)
            if largura >= altura:
                # PAISAGEM (Deitada) ou Quadrada: Trava a largura preenchendo a folha
                p_foto.add_run().add_picture(imagem_memoria, width=Inches(6.0))
            else:
                # RETRATO (Em pé): Trava a ALTURA para não ultrapassar o limite da folha A4
                p_foto.add_run().add_picture(imagem_memoria, height=Inches(7))
            
            # 3. CONTROLE DE RESPIRO
            p_respiro = doc.add_paragraph()
            p_respiro.paragraph_format.space_after = Pt(18)
            
        except Exception as e:
            print(f"❌ Erro ao processar a imagem {caminho_imagem}: {e}")
            
    return doc