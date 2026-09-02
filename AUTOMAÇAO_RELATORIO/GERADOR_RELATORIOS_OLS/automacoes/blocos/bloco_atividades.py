from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def pintar_fundo_celula(celula, cor_hex):
    """Injeta a formatação de cor de fundo direto no XML da célula."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_hex))
    celula._tc.get_or_add_tcPr().append(shading_elm)

def desenhar_atividades(doc, pacote_atividades):
    
    # 1. Dá um pequeno espaço antes de desenhar a barra
    doc.add_paragraph().paragraph_format.space_before = Pt(12)
    
    # 2. Cria a Barra Azul (Tabela 1x1)
    tabela_titulo = doc.add_table(rows=1, cols=1)
    celula = tabela_titulo.cell(0, 0)
    
    # Pinta o fundo de Azul OLS
    pintar_fundo_celula(celula, "17365D")
    
    # Adiciona e centraliza o texto "Atividade" em branco
    p_titulo = celula.paragraphs[0]
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_titulo = p_titulo.add_run("Atividade")
    run_titulo.bold = True
    run_titulo.font.size = Pt(12)
    run_titulo.font.color.rgb = RGBColor(255, 255, 255) # Fonte Branca
    
    # 3. Preenchimento dinâmico das atividades
    for item in pacote_atividades:
        data = item.get("data", "")
        desc = item.get("descricao", "")
        
        # Cria um parágrafo para a Data em Negrito (com um leve espaço em cima)
        p_data = doc.add_paragraph()
        p_data.paragraph_format.space_before = Pt(6)
        run_data = p_data.add_run(f"{data}:")
        run_data.bold = True
        
        # Cria o parágrafo para a Descrição (colado na data)
        p_desc = doc.add_paragraph(desc)
        p_desc.paragraph_format.space_after = Pt(6)

    return doc