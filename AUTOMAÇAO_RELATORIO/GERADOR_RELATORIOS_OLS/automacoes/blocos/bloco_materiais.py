from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def forcar_bordas(tabela):
    """Injeta a formatação da grade direto no XML, ignorando os estilos do Word."""
    tbl = tabela._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblBorders = OxmlElement('w:tblBorders')
    for borda in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        tag = OxmlElement(f'w:{borda}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')  # Espessura da linha
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), '000000') # Cor preta
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def pintar_fundo_celula(celula, cor_hex):
    """Injeta a formatação de cor de fundo direto no XML da célula."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_hex))
    celula._tc.get_or_add_tcPr().append(shading_elm)

def desenhar_materiais(doc, pacote_materiais):
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(18)
    run_titulo = p_titulo.add_run("Lista de Materiais / Equipamentos Utilizados")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    
    tabela = doc.add_table(rows=1, cols=5)
    forcar_bordas(tabela) 
    
    # 1. Definindo as larguras de cada coluna em Centímetros
    larguras = [Cm(1.5), Cm(7.0), Cm(3.5), Cm(2.5), Cm(2.5)]
    for i, largura in enumerate(larguras):
        tabela.columns[i].width = largura

    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Qtd'
    hdr_cells[1].text = 'Descrição'
    hdr_cells[2].text = 'Serial'
    hdr_cells[3].text = 'Asset Tag'
    hdr_cells[4].text = 'Ação'
    
    # 2. Pintando o cabeçalho (Fundo Azul, Letra Branca e Negrito)
    cor_azul_ols = "17365D" # Tom de azul escuro bem corporativo
    
    for i, cell in enumerate(hdr_cells):
        cell.width = larguras[i] # O Word precisa que force a largura na célula também
        pintar_fundo_celula(cell, cor_azul_ols)
        
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # Branco

    # =========================================================
    # 3. PREENCHENDO AS LINHAS DE DADOS (Este era o bloco que faltava!)
    # =========================================================
    for item in pacote_materiais:
        row_cells = tabela.add_row().cells
        row_cells[0].text = item.get("qtd", "")
        row_cells[1].text = item.get("descricao", "")
        row_cells[2].text = item.get("sn", "")     
        row_cells[3].text = item.get("asset", "")
        row_cells[4].text = item.get("acao", "")
        
        # Mantendo a largura correta para as linhas de baixo não entortarem
        for i, cell in enumerate(row_cells):
            cell.width = larguras[i]

    return doc